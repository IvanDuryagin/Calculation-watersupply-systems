#interface.py

import csv
import tkinter as tk
from tkinter import messagebox, StringVar, Toplevel, ttk
import numpy as np
import data
from openpyxl import Workbook
from docx import Document
import tkinter.filedialog
from typing import List, Tuple, Dict, Optional, Union, Any
from dataclasses import dataclass
from enum import Enum

@dataclass
class CalculationResult:
    t: float
    U: float
    D: int
    q_chru: float
    q_c0: float
    group: int
    consumer: str
    PcN: float
    alpha: float
    Q: float
    velocity: float

class FileType(Enum):
    CSV = ("CSV files", "*.csv")
    EXCEL = ("Excel files", "*.xlsx")
    WORD = ("Word documents", "*.docx")

class ConsumerCalculator:
    RESULT_HEADERS = [
        "t", "U", "D", "q_(c)hru", "q_(c)0", 
        "Группа (t_num)", "Потребитель", 
        "Pc*N", "alpha", "Q", "Скорость"
    ]
    
    TEXT = {
        "title": "Гидравлический расчёт системы водоснабжения по СП 30.13330.2020",
        "consumer_label": "СП 30.13330",
        "t_entry_default": "Введите номер потребителя",
        "u_entry_default": "Введите U",
        "add_section": "➕ Добавить участок",
        "delete_section": "❌ Удалить участок",
        "calculate": "▶️ Рассчитать",
        "show_results": "↪️ Результаты расчёта",
        "save_results": "⏫ Сохранить результаты",
        "success": "Успех",
        "error": "Ошибка",
        "file_not_saved": "Файл не сохранён.",
        "calculation_success": "Расчёт выполнен успешно.",
        "no_data": "Нет данных для отображения.",
        "invalid_t": "Введите тип потребителя.",
        "t_not_found": "Значение t={} не найдено в массиве.",
        "section_error": "Ошибка в расчётах участка: {}",
        "interpolation_error": "Не удалось выполнить интерполяцию.",
        "velocity_error": "Не удалось вычислить значение скорости."
    }

    def __init__(self, master: tk.Tk):
        self.master = master
        self.current_consumer_data = None
        self.master.title(self.TEXT["title"])
        self.master.minsize(800, 600)
        
        self.entries = []
        self.results = []
        self.current_consumer = ""
        
        self._setup_ui()
        self.add_section()

    def _setup_ui(self) -> None:
        self._create_input_panel()
        self._create_scrollable_area()
      
    def _create_input_panel(self) -> None:
        input_frame = ttk.Frame(self.master, padding="10")
        input_frame.pack(fill=tk.X)
        
        ttk.Label(input_frame, text=self.TEXT["consumer_label"]).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            input_frame, text="◀️ Расчёт нагрузок",
            command=self.open_load_calculator
        ).pack(side=tk.LEFT, padx=5)
        
        self.t_entry = ttk.Entry(input_frame)
        self.t_entry.pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        self.t_entry.insert(0, self.TEXT["t_entry_default"])
        
        buttons = [
            (self.TEXT["calculate"], self.calculate),
            (self.TEXT["show_results"], self.show_results),
            (self.TEXT["save_results"], self.save_results)
        ]
        
        for text, command in buttons:
            ttk.Button(
                input_frame, text=text, command=command
            ).pack(side=tk.LEFT, padx=5)

    def open_load_calculator(self):
        try:
            if not self.current_consumer_data:
                raise ValueError("Сначала выполните гидравлический расчёт")
                
            LoadCalculator(self.master, self.current_consumer_data)
        except ValueError as e:
            messagebox.showerror("Ошибка", str(e))

    def _create_scrollable_area(self) -> None:
        container = ttk.Frame(self.master)
        container.pack(fill=tk.BOTH, expand=True)
        
        self.canvas = tk.Canvas(container)
        scrollbar_y = ttk.Scrollbar(container, orient="vertical", command=self.canvas.yview)
        scrollbar_x = ttk.Scrollbar(container, orient="horizontal", command=self.canvas.xview)
        
        self.scrollable_frame = ttk.Frame(self.canvas)
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=scrollbar_y.set)
        self.canvas.configure(xscrollcommand=scrollbar_x.set)
        
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)

    def add_section(self, after_frame: Optional[tk.Frame] = None) -> None:
        new_frame = ttk.Frame(self.scrollable_frame, padding="5")
    
        u_entry = ttk.Entry(new_frame)
        u_entry.pack(side=tk.LEFT, padx=5)
        u_entry.insert(0, self.TEXT["u_entry_default"])

        diam_values = [d for d in data.diam_values if d > 0]
        diam_str_values = [str(int(d)) for d in diam_values]
    
        diam_var = StringVar(value=diam_str_values[0])
        diam_menu = ttk.OptionMenu(
            new_frame, 
            diam_var, 
            diam_str_values[0],
            *diam_str_values
        )
        diam_menu.pack(side=tk.LEFT, padx=5)

        ttk.Button(
            new_frame, text=self.TEXT["add_section"],
            command=lambda: self.add_section(new_frame)
        ).pack(side=tk.LEFT, padx=5)
    
        ttk.Button(
            new_frame, text=self.TEXT["delete_section"],
            command=lambda: self.remove_section(new_frame)
        ).pack(side=tk.LEFT, padx=5)

        if after_frame:
            for i, (_, _, frame) in enumerate(self.entries):
                if frame == after_frame:
                    self.entries.insert(i+1, (u_entry, diam_var, new_frame))
                    new_frame.pack(in_=self.scrollable_frame, after=after_frame)
                    break
        else:
            self.entries.append((u_entry, diam_var, new_frame))
            new_frame.pack(in_=self.scrollable_frame)

    def remove_section(self, frame: tk.Frame) -> None:
        if len(self.entries) <= 1:
            messagebox.showwarning(self.TEXT["error"], "Нельзя удалить последний участок")
            return
            
        frame.destroy()
        self.entries = [entry for entry in self.entries if entry[2] != frame]

    def calculate(self) -> None:
        self.results.clear()
        
        try:
            t_input = self._validate_t_input()
            index, params = self._get_consumer_params(t_input)
            
            self.current_consumer_data = {
                't_input': t_input,
                'index': index,
                'params': params,
                'q_tot_hru': data.q_tot_hru_values[index],
                'q_h_hru': data.q_h_hru_values[index],
                'q_c_hru': data.q_c_hru_values[index],
                'q_tot': data.q_tot_values[index],
                'q_h': data.q_h_values[index],
                'q_c': data.q_c_values[index],
                'q_tot_0': data.q_tot_0_values[index],
                'q_h_0': data.q_h_0_values[index],
                'q_c_0': data.q_c_0_values[index],
                'q_tot_0_hr': data.q_tot_0_hr_values[index],
                'q_h_0_hr': data.q_h_0_hr_values[index],
                'q_c_0_hr': data.q_c_0_hr_values[index],
                'consumer_name': params[3]
            }
            self.current_consumer = params[3]
            
            for u_entry, diam_var, _ in self.entries:
                result = self._calculate_section(u_entry, diam_var, t_input, *params)
                self.results.append(result)
                
            messagebox.showinfo(self.TEXT["success"], self.TEXT["calculation_success"])
        except ValueError as e:
            messagebox.showerror(self.TEXT["error"], str(e))

    def _validate_t_input(self) -> float:
        try:
            t_input = float(self.t_entry.get())
            if t_input not in data.t_values:
                raise ValueError(self.TEXT["t_not_found"].format(t_input))
            return t_input
        except ValueError:
            raise ValueError(self.TEXT["invalid_t"])

    def _get_consumer_params(self, t_input: float) -> Tuple[int, Tuple[float, float, int, str]]:
        index = np.where(data.t_values == t_input)[0][0]
        return index, (
            data.q_c_hru_values[index],
            data.q_c_0_values[index],
            data.t_num_values[index],
            data.t_string_values[index]
        )

    def _calculate_section(self, u_entry: tk.Entry, diam_var: StringVar, 
                          t_input: float, q_c_hru: float, q_c_0: float, 
                          t_num: int, t_string: str) -> CalculationResult:
        try:
            U = float(u_entry.get())
            D = int(diam_var.get())
            v_values = data.v_dict[D]

            x_input = (q_c_hru * U) / (3600 * q_c_0)
            alpha = self._interpolate(x_input, data.x_values, data.y_values)
            Q = 5 * q_c_0 * alpha
            velocity = self._interpolate_velocity(Q, v_values)

            return CalculationResult(
                t=t_input, U=U, D=D, q_chru=q_c_hru, q_c0=q_c_0,
                group=t_num, consumer=t_string, PcN=x_input,
                alpha=alpha, Q=Q, velocity=velocity
            )
        except ValueError as e:
            raise ValueError(self.TEXT["section_error"].format(str(e)))

    def _interpolate_velocity(self, Q: float, v_values: List[float]) -> float:
        q_values = data.q_values
        try:
            if Q < q_values[0]:
                return self._linear_interpolation(Q, q_values[0], q_values[1], v_values[0], v_values[1])
            if Q > q_values[-1]:
                return self._linear_interpolation(Q, q_values[-2], q_values[-1], v_values[-2], v_values[-1])
            
            for i in range(len(q_values) - 1):
                if q_values[i] <= Q <= q_values[i + 1]:
                    return self._linear_interpolation(Q, q_values[i], q_values[i + 1], v_values[i], v_values[i + 1])
        except IndexError:
            pass
            
        raise ValueError(self.TEXT["velocity_error"])

    @staticmethod
    def _interpolate(x: float, x_values: List[float], y_values: List[float]) -> float:
        try:
            if x < x_values[0]:
                return ConsumerCalculator._linear_interpolation(x, x_values[0], x_values[1], y_values[0], y_values[1])
            if x > x_values[-1]:
                return ConsumerCalculator._linear_interpolation(x, x_values[-2], x_values[-1], y_values[-2], y_values[-1])
            
            for i in range(len(x_values) - 1):
                if x_values[i] <= x <= x_values[i + 1]:
                    return ConsumerCalculator._linear_interpolation(x, x_values[i], x_values[i + 1], y_values[i], y_values[i + 1])
        except IndexError:
            pass
            
        raise ValueError(ConsumerCalculator.TEXT["interpolation_error"])

    @staticmethod
    def _linear_interpolation(x: float, x0: float, x1: float, y0: float, y1: float) -> float:
        return y0 + (y1 - y0) * (x - x0) / (x1 - x0)

    def show_results(self) -> None:
        if not self.results:
            messagebox.showwarning(self.TEXT["error"], self.TEXT["no_data"])
            return
        
        results_window = Toplevel(self.master)
        results_window.title(self.TEXT["show_results"])
        results_window.minsize(800, 400)
        
        ttk.Label(
            results_window, 
            text=f"Потребитель: {self.current_consumer}", 
            font=("Arial", 12, "bold")
        ).pack(pady=10)

        tree = ttk.Treeview(results_window, columns=self.RESULT_HEADERS, show="headings")
        
        for header in self.RESULT_HEADERS:
            tree.heading(header, text=header)
            tree.column(header, width=100, anchor=tk.CENTER)
        
        for result in self.results:
            tree.insert("", tk.END, values=[
                result.t, result.U, result.D, result.q_chru, result.q_c0,
                result.group, result.consumer, round(result.PcN, 4),
                round(result.alpha, 4), round(result.Q, 4), round(result.velocity, 4)
            ])
        
        y_scroll = ttk.Scrollbar(results_window, orient="vertical", command=tree.yview)
        x_scroll = ttk.Scrollbar(results_window, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=y_scroll.set, xscrollcommand=x_scroll.set)
        
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        y_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        x_scroll.pack(side=tk.BOTTOM, fill=tk.X)

    def save_results(self) -> None:
        if not self.results:
            messagebox.showwarning(self.TEXT["error"], self.TEXT["no_data"])
            return
        
        file_path = tk.filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[ft.value for ft in FileType],
            title="Сохранить результаты"
        )
        
        if not file_path:
            messagebox.showwarning(self.TEXT["file_not_saved"], self.TEXT["file_not_saved"])
            return
        
        try:
            extension = file_path.split('.')[-1].lower()
            if extension == 'csv':
                self._save_to_csv(file_path)
            elif extension == 'xlsx':
                self._save_to_excel(file_path)
            elif extension == 'docx':
                self._save_to_docx(file_path)
            else:
                messagebox.showerror(self.TEXT["error"], "Неподдерживаемый формат файла")
                return
                
            messagebox.showinfo(self.TEXT["success"], f"Файл успешно сохранён: {file_path}")
        except Exception as e:
            messagebox.showerror(self.TEXT["error"], f"Ошибка при сохранении: {str(e)}")

    def _save_to_csv(self, file_path: str) -> None:
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(self.RESULT_HEADERS)
            for result in self.results:
                writer.writerow([
                    result.t, result.U, result.D, result.q_chru, result.q_c0,
                    result.group, result.consumer, result.PcN,
                    result.alpha, result.Q, result.velocity
                ])

    def _save_to_excel(self, file_path: str) -> None:
        wb = Workbook()
        ws = wb.active
        ws.title = "Результаты расчёта"
        ws.append(self.RESULT_HEADERS)
        
        for result in self.results:
            ws.append([
                result.t, result.U, result.D, result.q_chru, result.q_c0,
                result.group, result.consumer, result.PcN,
                result.alpha, result.Q, result.velocity
            ])
            
        wb.save(file_path)

    def _save_to_docx(self, file_path: str) -> None:
        doc = Document()
        doc.add_heading(f"Результаты расчёта: {self.current_consumer}", level=1)
        
        table = doc.add_table(rows=1, cols=len(self.RESULT_HEADERS))
        table.style = 'Table Grid'
        
        for i, header in enumerate(self.RESULT_HEADERS):
            table.cell(0, i).text = header
            
        for result in self.results:
            row = table.add_row().cells
            values = [
                str(result.t), str(result.U), str(result.D), str(result.q_chru),
                str(result.q_c0), str(result.group), result.consumer,
                f"{result.PcN:.4f}", f"{result.alpha:.4f}", 
                f"{result.Q:.4f}", f"{result.velocity:.4f}"
            ]
            for i, value in enumerate(values):
                row[i].text = value
                
        doc.save(file_path)
        
    def load_imported_data(self, data):
        """Загружает импортированные данные в интерфейс"""
        # Сначала очищаем существующие участки
        while len(self.entries) > 1:
            self.remove_section(self.entries[0][2])
    
        # Добавляем участки из импортированных данных
        for item in data:
            self.add_section()
            last_idx = len(self.entries) - 1
            self.entries[last_idx][0].delete(0, tk.END)  # очищаем поле U
            self.entries[last_idx][0].insert(0, str(item['U']))  # заполняем U
        
            # Устанавливаем диаметр, если он есть в данных
            if 'D' in item:
                self.entries[last_idx][1].set(str(int(item['D'])))
        
            # Устанавливаем номер участка (t), если есть поле ввода
            if hasattr(self, 't_entry'):
                self.t_entry.delete(0, tk.END)
                self.t_entry.insert(0, str(item.get('t', '')))
    
    def validate_import_data(self, data):
        """Проверяет корректность импортированных данных"""
        validated = []
        for item in data:
            try:
                validated.append({
                    't': float(item['t']),
                    'U': float(item['U']),
                    'D': int(item['D'])
                })
            except (ValueError, KeyError):
                continue
        return validated

class HotWaterCalculator:
    RESULT_HEADERS = [
        "t", "U", "D", "q_(h)hru", "q_(h)0", 
        "Группа (t_num)", "Потребитель", 
        "Ph*N", "alpha", "Q", "Скорость"
    ]
    
    TEXT = {
        "title": "Гидравлический расчёт системы горячего водоснабжения по СП 30.13330.2020",
        "consumer_label": "СП 30.13330",
        "t_entry_default": "Введите номер потребителя",
        "u_entry_default": "Введите U",
        "add_section": "➕ Добавить участок",
        "delete_section": "❌ Удалить участок",
        "calculate": "▶️ Рассчитать",
        "show_results": "↪️ Результаты расчёта",
        "save_results": "⏫ Сохранить результаты",
        "success": "Успех",
        "error": "Ошибка",
        "file_not_saved": "Файл не сохранён.",
        "calculation_success": "Расчёт выполнен успешно.",
        "no_data": "Нет данных для отображения.",
        "invalid_t": "Введите корректное значение для t.",
        "t_not_found": "Значение t={} не найдено в массиве.",
        "section_error": "Ошибка в расчётах участка: {}",
        "interpolation_error": "Не удалось выполнить интерполяцию.",
        "velocity_error": "Не удалось вычислить значение скорости."
    }

    def __init__(self, master: tk.Tk):
        self.master = master
        self.current_consumer_data = None
        self.master.title(self.TEXT["title"])
        self.master.minsize(800, 600)
        
        self.entries = []
        self.results = []
        self.current_consumer = ""
        
        self._setup_ui()
        self.add_section()

    def _setup_ui(self) -> None:
        self._create_input_panel()
        self._create_scrollable_area()
      
    def _create_input_panel(self) -> None:
        input_frame = ttk.Frame(self.master, padding="10")
        input_frame.pack(fill=tk.X)
        
        ttk.Label(input_frame, text=self.TEXT["consumer_label"]).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            input_frame, text="◀️ Расчёт нагрузок",
            command=self.open_load_calculator
        ).pack(side=tk.LEFT, padx=5)
        
        self.t_entry = ttk.Entry(input_frame)
        self.t_entry.pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        self.t_entry.insert(0, self.TEXT["t_entry_default"])
        
        buttons = [
            (self.TEXT["calculate"], self.calculate),
            (self.TEXT["show_results"], self.show_results),
            (self.TEXT["save_results"], self.save_results)
        ]
        
        for text, command in buttons:
            ttk.Button(
                input_frame, text=text, command=command
            ).pack(side=tk.LEFT, padx=5)

    def open_load_calculator(self):
        try:
            if not self.current_consumer_data:
                raise ValueError("Сначала выполните расчёт потребителя")
                
            LoadCalculator(self.master, self.current_consumer_data)
        except ValueError as e:
            messagebox.showerror("Ошибка", str(e))

    def _create_scrollable_area(self) -> None:
        container = ttk.Frame(self.master)
        container.pack(fill=tk.BOTH, expand=True)
        
        self.canvas = tk.Canvas(container)
        scrollbar_y = ttk.Scrollbar(container, orient="vertical", command=self.canvas.yview)
        scrollbar_x = ttk.Scrollbar(container, orient="horizontal", command=self.canvas.xview)
        
        self.scrollable_frame = ttk.Frame(self.canvas)
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=scrollbar_y.set)
        self.canvas.configure(xscrollcommand=scrollbar_x.set)
        
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)

    def add_section(self, after_frame: Optional[tk.Frame] = None) -> None:
        new_frame = ttk.Frame(self.scrollable_frame, padding="5")
    
        u_entry = ttk.Entry(new_frame)
        u_entry.pack(side=tk.LEFT, padx=5)
        u_entry.insert(0, self.TEXT["u_entry_default"])

        diam_values = [d for d in data.diam_values if d > 0]
        diam_str_values = [str(int(d)) for d in diam_values]
    
        diam_var = StringVar(value=diam_str_values[0])
        diam_menu = ttk.OptionMenu(
            new_frame, 
            diam_var, 
            diam_str_values[0],
            *diam_str_values
        )
        diam_menu.pack(side=tk.LEFT, padx=5)

        ttk.Button(
            new_frame, text=self.TEXT["add_section"],
            command=lambda: self.add_section(new_frame)
        ).pack(side=tk.LEFT, padx=5)
    
        ttk.Button(
            new_frame, text=self.TEXT["delete_section"],
            command=lambda: self.remove_section(new_frame)
        ).pack(side=tk.LEFT, padx=5)

        if after_frame:
            for i, (_, _, frame) in enumerate(self.entries):
                if frame == after_frame:
                    self.entries.insert(i+1, (u_entry, diam_var, new_frame))
                    new_frame.pack(in_=self.scrollable_frame, after=after_frame)
                    break
        else:
            self.entries.append((u_entry, diam_var, new_frame))
            new_frame.pack(in_=self.scrollable_frame)

    def remove_section(self, frame: tk.Frame) -> None:
        if len(self.entries) <= 1:
            messagebox.showwarning(self.TEXT["error"], "Нельзя удалить последний участок")
            return
            
        frame.destroy()
        self.entries = [entry for entry in self.entries if entry[2] != frame]

    def calculate(self) -> None:
        self.results.clear()
        
        try:
            t_input = self._validate_t_input()
            index, params = self._get_consumer_params(t_input)
            
            self.current_consumer_data = {
                't_input': t_input,
                'index': index,
                'params': params,
                'q_tot_hru': data.q_tot_hru_values[index],
                'q_h_hru': data.q_h_hru_values[index],
                'q_c_hru': data.q_c_hru_values[index],
                'q_tot': data.q_tot_values[index],
                'q_h': data.q_h_values[index],
                'q_c': data.q_c_values[index],
                'q_tot_0': data.q_tot_0_values[index],
                'q_h_0': data.q_h_0_values[index],
                'q_c_0': data.q_c_0_values[index],
                'q_tot_0_hr': data.q_tot_0_hr_values[index],
                'q_h_0_hr': data.q_h_0_hr_values[index],
                'q_c_0_hr': data.q_c_0_hr_values[index],
                'consumer_name': params[3]
            }
            self.current_consumer = params[3]
            
            for u_entry, diam_var, _ in self.entries:
                result = self._calculate_section(u_entry, diam_var, t_input, *params)
                self.results.append(result)
                
            messagebox.showinfo(self.TEXT["success"], self.TEXT["calculation_success"])
        except ValueError as e:
            messagebox.showerror(self.TEXT["error"], str(e))

    def _validate_t_input(self) -> float:
        try:
            t_input = float(self.t_entry.get())
            if t_input not in data.t_values:
                raise ValueError(self.TEXT["t_not_found"].format(t_input))
            return t_input
        except ValueError:
            raise ValueError(self.TEXT["invalid_t"])

    def _get_consumer_params(self, t_input: float) -> Tuple[int, Tuple[float, float, int, str]]:
        index = np.where(data.t_values == t_input)[0][0]
        return index, (
            data.q_h_hru_values[index],  # Используем q_h_hru вместо q_c_hru
            data.q_h_0_values[index],    # Используем q_h_0 вместо q_c_0
            data.t_num_values[index],
            data.t_string_values[index]
        )

    def _calculate_section(self, u_entry: tk.Entry, diam_var: StringVar, 
                          t_input: float, q_h_hru: float, q_h_0: float, 
                          t_num: int, t_string: str) -> CalculationResult:
        try:
            U = float(u_entry.get())
            D = int(diam_var.get())
            v_values = data.v_dict[D]

            x_input = (q_h_hru * U) / (3600 * q_h_0)  # Используем параметры горячей воды
            alpha = self._interpolate(x_input, data.x_values, data.y_values)
            Q = 5 * q_h_0 * alpha  # Используем q_h_0 вместо q_c_0
            velocity = self._interpolate_velocity(Q, v_values)

            return CalculationResult(
                t=t_input, U=U, D=D, 
                q_chru=q_h_hru,  # Используем q_h_hru вместо q_c_hru
                q_c0=q_h_0,      # Используем q_h_0 вместо q_c_0
                group=t_num, 
                consumer=t_string, 
                PcN=x_input,
                alpha=alpha, 
                Q=Q, 
                velocity=velocity
            )
        except ValueError as e:
            raise ValueError(self.TEXT["section_error"].format(str(e)))

    def _interpolate_velocity(self, Q: float, v_values: List[float]) -> float:
        q_values = data.q_values
        try:
            if Q < q_values[0]:
                return self._linear_interpolation(Q, q_values[0], q_values[1], v_values[0], v_values[1])
            if Q > q_values[-1]:
                return self._linear_interpolation(Q, q_values[-2], q_values[-1], v_values[-2], v_values[-1])
            
            for i in range(len(q_values) - 1):
                if q_values[i] <= Q <= q_values[i + 1]:
                    return self._linear_interpolation(Q, q_values[i], q_values[i + 1], v_values[i], v_values[i + 1])
        except IndexError:
            pass
            
        raise ValueError(self.TEXT["velocity_error"])

    @staticmethod
    def _interpolate(x: float, x_values: List[float], y_values: List[float]) -> float:
        try:
            if x < x_values[0]:
                return HotWaterCalculator._linear_interpolation(x, x_values[0], x_values[1], y_values[0], y_values[1])
            if x > x_values[-1]:
                return HotWaterCalculator._linear_interpolation(x, x_values[-2], x_values[-1], y_values[-2], y_values[-1])
            
            for i in range(len(x_values) - 1):
                if x_values[i] <= x <= x_values[i + 1]:
                    return HotWaterCalculator._linear_interpolation(x, x_values[i], x_values[i + 1], y_values[i], y_values[i + 1])
        except IndexError:
            pass
            
        raise ValueError(HotWaterCalculator.TEXT["interpolation_error"])

    @staticmethod
    def _linear_interpolation(x: float, x0: float, x1: float, y0: float, y1: float) -> float:
        return y0 + (y1 - y0) * (x - x0) / (x1 - x0)

    def show_results(self) -> None:
        if not self.results:
            messagebox.showwarning(self.TEXT["error"], self.TEXT["no_data"])
            return
        
        results_window = Toplevel(self.master)
        results_window.title(self.TEXT["show_results"])
        results_window.minsize(800, 400)
        
        ttk.Label(
            results_window, 
            text=f"Потребитель: {self.current_consumer}", 
            font=("Arial", 12, "bold")
        ).pack(pady=10)

        tree = ttk.Treeview(results_window, columns=self.RESULT_HEADERS, show="headings")
        
        for header in self.RESULT_HEADERS:
            tree.heading(header, text=header)
            tree.column(header, width=100, anchor=tk.CENTER)
        
        for result in self.results:
            tree.insert("", tk.END, values=[
                result.t, result.U, result.D, result.q_chru, result.q_c0,
                result.group, result.consumer, round(result.PcN, 4),
                round(result.alpha, 4), round(result.Q, 4), round(result.velocity, 4)
            ])
        
        y_scroll = ttk.Scrollbar(results_window, orient="vertical", command=tree.yview)
        x_scroll = ttk.Scrollbar(results_window, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=y_scroll.set, xscrollcommand=x_scroll.set)
        
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        y_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        x_scroll.pack(side=tk.BOTTOM, fill=tk.X)

    def save_results(self) -> None:
        if not self.results:
            messagebox.showwarning(self.TEXT["error"], self.TEXT["no_data"])
            return
        
        file_path = tk.filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[ft.value for ft in FileType],
            title="Сохранить результаты"
        )
        
        if not file_path:
            messagebox.showwarning(self.TEXT["file_not_saved"], self.TEXT["file_not_saved"])
            return
        
        try:
            extension = file_path.split('.')[-1].lower()
            if extension == 'csv':
                self._save_to_csv(file_path)
            elif extension == 'xlsx':
                self._save_to_excel(file_path)
            elif extension == 'docx':
                self._save_to_docx(file_path)
            else:
                messagebox.showerror(self.TEXT["error"], "Неподдерживаемый формат файла")
                return
                
            messagebox.showinfo(self.TEXT["success"], f"Файл успешно сохранён: {file_path}")
        except Exception as e:
            messagebox.showerror(self.TEXT["error"], f"Ошибка при сохранении: {str(e)}")

    def _save_to_csv(self, file_path: str) -> None:
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(self.RESULT_HEADERS)
            for result in self.results:
                writer.writerow([
                    result.t, result.U, result.D, result.q_chru, result.q_c0,
                    result.group, result.consumer, result.PcN,
                    result.alpha, result.Q, result.velocity
                ])

    def _save_to_excel(self, file_path: str) -> None:
        wb = Workbook()
        ws = wb.active
        ws.title = "Результаты расчёта"
        ws.append(self.RESULT_HEADERS)
        
        for result in self.results:
            ws.append([
                result.t, result.U, result.D, result.q_chru, result.q_c0,
                result.group, result.consumer, result.PcN,
                result.alpha, result.Q, result.velocity
            ])
            
        wb.save(file_path)

    def _save_to_docx(self, file_path: str) -> None:
        doc = Document()
        doc.add_heading(f"Результаты расчёта: {self.current_consumer}", level=1)
        
        table = doc.add_table(rows=1, cols=len(self.RESULT_HEADERS))
        table.style = 'Table Grid'
        
        for i, header in enumerate(self.RESULT_HEADERS):
            table.cell(0, i).text = header
            
        for result in self.results:
            row = table.add_row().cells
            values = [
                str(result.t), str(result.U), str(result.D), str(result.q_chru),
                str(result.q_c0), str(result.group), result.consumer,
                f"{result.PcN:.4f}", f"{result.alpha:.4f}", 
                f"{result.Q:.4f}", f"{result.velocity:.4f}"
            ]
            for i, value in enumerate(values):
                row[i].text = value
                
        doc.save(file_path)

    def load_imported_data(self, data):
        """Загружает импортированные данные в интерфейс"""
        # Сначала очищаем существующие участки
        while len(self.entries) > 1:
            self.remove_section(self.entries[0][2])
    
        # Добавляем участки из импортированных данных
        for item in data:
            self.add_section()
            last_idx = len(self.entries) - 1
            self.entries[last_idx][0].delete(0, tk.END)  # очищаем поле U
            self.entries[last_idx][0].insert(0, str(item['U']))  # заполняем U
        
            # Устанавливаем диаметр, если он есть в данных
            if 'D' in item:
                self.entries[last_idx][1].set(str(int(item['D'])))
        
            # Устанавливаем номер участка (t), если есть поле ввода
            if hasattr(self, 't_entry'):
                self.t_entry.delete(0, tk.END)
                self.t_entry.insert(0, str(item.get('t', '')))
    
    def validate_import_data(self, data):
        """Проверяет корректность импортированных данных"""
        validated = []
        for item in data:
            try:
                validated.append({
                    't': float(item['t']),
                    'U': float(item['U']),
                    'D': int(item['D'])
                })
            except (ValueError, KeyError):
                continue
        return validated

class LoadCalculator:
    TEXT = {
        "title": "◀️ Расчёт нагрузок",
        "u_label": "Общее число потребителей (U):",
        "calculate": "▶️ Рассчитать",
        "save_results": "⏫ Сохранить результаты",
        "results_title": "Результаты расчёта нагрузок"
    }

    def __init__(self, master, consumer_data):
        self.top = tk.Toplevel(master)
        self.consumer_data = consumer_data
        self.top.title(self.TEXT["title"])
        self.top.geometry("900x600")
        
        self._setup_ui()
        
    def _setup_ui(self):
        input_frame = ttk.Frame(self.top, padding="10")
        input_frame.pack(fill=tk.X)
        
        ttk.Label(input_frame, text=self.TEXT["u_label"]).pack(side=tk.LEFT)
        self.u_entry = ttk.Entry(input_frame, width=15)
        self.u_entry.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            input_frame, text=self.TEXT["calculate"],
            command=self.calculate
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            input_frame, text=self.TEXT["save_results"],
            command=self.save_results
        ).pack(side=tk.LEFT)
        
        self.results_frame = ttk.Frame(self.top, padding="10")
        self.results_frame.pack(fill=tk.BOTH, expand=True)
        
        self.tree = ttk.Treeview(self.results_frame, columns=("Parameter", "Value", "Unit"), show="headings")
        self.tree.heading("Parameter", text="Параметр")
        self.tree.heading("Value", text="Значение")
        self.tree.heading("Unit", text="Ед. изм.")
        self.tree.column("Parameter", width=300)
        self.tree.column("Value", width=200)
        self.tree.column("Unit", width=100)
        
        scroll_y = ttk.Scrollbar(self.results_frame, orient="vertical", command=self.tree.yview)
        scroll_x = ttk.Scrollbar(self.results_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        scroll_x.pack(side=tk.BOTTOM, fill=tk.X)
        
        temp_frame = ttk.Frame(self.top, padding="10")
        temp_frame.pack(fill=tk.X)
        
        ttk.Label(temp_frame, text="Температура горячей воды (t_h, oC):").pack(side=tk.LEFT)
        self.t_h_entry = ttk.Entry(temp_frame, width=5)
        self.t_h_entry.insert(0, "60")
        self.t_h_entry.pack(side=tk.LEFT, padx=5)
        
        ttk.Label(temp_frame, text="Температура холодной воды (t_c, oC):").pack(side=tk.LEFT, padx=(10,0))
        self.t_c_entry = ttk.Entry(temp_frame, width=5)
        self.t_c_entry.insert(0, "5")
        self.t_c_entry.pack(side=tk.LEFT)

        ttk.Label(temp_frame, text="Период водопотребления (T, ч):").pack(side=tk.LEFT, padx=(10,0))
        self.T_entry = ttk.Entry(temp_frame, width=5)
        self.T_entry.insert(0, "24")
        self.T_entry.pack(side=tk.LEFT)
    
    def calculate(self):
        try:
            U = float(self.u_entry.get())
            t_h = float(self.t_h_entry.get())
            t_c = float(self.t_c_entry.get())
            T = float(self.T_entry.get())
            
            for item in self.tree.get_children():
                self.tree.delete(item)
                
            if self.consumer_data:
                idx = self.consumer_data['index']
                
                results = [
                    ("Секундная вероятность действия приборов обшая (P_tot*N)", 
                     (self.consumer_data['q_tot_hru'] * U) / (3600 * self.consumer_data['q_tot_0']), 
                     "-"),
                    ("Секундная вероятность действия приборов на ГВС (P_h*N)", 
                     (self.consumer_data['q_h_hru'] * U) / (3600 * self.consumer_data['q_h_0']), 
                     "-"),
                    ("Секундная вероятность действия приборов на ХВС (P_c*N)", 
                     (self.consumer_data['q_c_hru'] * U) / (3600 * self.consumer_data['q_c_0']), 
                     "-"),

                    ("Альфа секундная общая (a_tot)", 
                     self._calculate_alpha(self.consumer_data['q_tot_hru'], self.consumer_data['q_tot_0'], U), 
                     "-"),
                    ("Альфа секундная на ГВС (a_h)", 
                     self._calculate_alpha(self.consumer_data['q_h_hru'], self.consumer_data['q_h_0'], U), 
                     "-"),
                    ("Альфа секундная на ХВС (a_c)", 
                     self._calculate_alpha(self.consumer_data['q_c_hru'], self.consumer_data['q_c_0'], U), 
                     "-"),

                    ("Расчётный секундный расход общий (q_tot)", 
                     5 * self.consumer_data['q_tot_0'] * self._calculate_alpha(self.consumer_data['q_tot_hru'], self.consumer_data['q_tot_0'], U), 
                     "л/с"),
                    ("Расчётный секундный расход на ГВС (q_h)", 
                     5 * self.consumer_data['q_h_0'] * self._calculate_alpha(self.consumer_data['q_h_hru'], self.consumer_data['q_h_0'], U), 
                     "л/с"),
                    ("Расчётный секундный расход на ХВС (q_c)", 
                     5 * self.consumer_data['q_c_0'] * self._calculate_alpha(self.consumer_data['q_c_hru'], self.consumer_data['q_c_0'], U), 
                     "л/с"),
                     
                    ("Часовая вероятность действия приборов общая (P_tot*N)", 
                     self.consumer_data['q_tot_hru'] * U / self.consumer_data['q_tot_0_hr'], 
                     "-"),
                    ("Часовая вероятность действия приборов на ГВС(P_h*N)", 
                     self.consumer_data['q_h_hru'] * U / self.consumer_data['q_h_0_hr'], 
                     "-"),
                    ("Часовая вероятность действия приборов на ХВС (P_c*N)", 
                     self.consumer_data['q_c_hru'] * U / self.consumer_data['q_c_0_hr'], 
                     "-"),

                    ("Альфа часовая общая (a_tot_hr)", 
                     self._calculate_alpha_h(self.consumer_data['q_tot_hru'], self.consumer_data['q_tot_0_hr'], U), 
                     "-"),
                    ("Альфа часовая на ГВС(a_h_hr)", 
                     self._calculate_alpha_h(self.consumer_data['q_h_hru'], self.consumer_data['q_h_0_hr'], U), 
                     "-"),
                    ("Альфа часовая на ХВС (a_c_hr)", 
                     self._calculate_alpha_h(self.consumer_data['q_c_hru'], self.consumer_data['q_c_0_hr'], U), 
                     "-"),
                     
                    ("Часовой расход общий (q_tot_hr)", 
                     0.005 * self.consumer_data['q_tot_0_hr'] * self._calculate_alpha_h(self.consumer_data['q_tot_hru'], self.consumer_data['q_tot_0_hr'], U), 
                     "м³/ч"),
                    ("Часовой расход на ГВС (q_h_hr)", 
                     0.005 * self.consumer_data['q_h_0_hr'] * self._calculate_alpha_h(self.consumer_data['q_h_hru'], self.consumer_data['q_h_0_hr'], U), 
                     "м³/ч"),
                    ("Часовой расход на ХВС (q_c_hr)", 
                     0.005 * self.consumer_data['q_c_0_hr'] * self._calculate_alpha_h(self.consumer_data['q_c_hru'], self.consumer_data['q_c_0_hr'], U), 
                     "м³/ч"),
                     
                    ("Суточный расход общий (Q_сут_tot)", 
                     self.consumer_data['q_tot'] * U / 1000, 
                     "м³/сут"),
                    ("Суточный расход на ГВС (Q_сут_h)", 
                     self.consumer_data['q_h'] * U / 1000, 
                     "м³/сут"),
                    ("Суточный расход на ХВС (Q_сут_c)", 
                     self.consumer_data['q_c'] * U / 1000, 
                     "м³/сут"),
                     
                    ("Расход тепла на ГВС максимальный (Q(h,hr))", 
                     1.16 * (0.005 * self.consumer_data['q_h_0_hr'] * self._calculate_alpha_h(self.consumer_data['q_h_hru'], self.consumer_data['q_h_0_hr'], U)) * (t_h - t_c) + 0.4 * (0.005 * self.consumer_data['q_h_0_hr'] * self._calculate_alpha_h(self.consumer_data['q_h_hru'], self.consumer_data['q_h_0_hr'], U)), 
                     "кВт"),

                    ("Расход тепла на ГВС максимальный (Q(h,hr))", 
                     0.0008598452*(1.16 * (0.005 * self.consumer_data['q_h_0_hr'] * self._calculate_alpha_h(self.consumer_data['q_h_hru'], self.consumer_data['q_h_0_hr'], U)) * (t_h - t_c) + 0.4 * (0.005 * self.consumer_data['q_h_0_hr'] * self._calculate_alpha_h(self.consumer_data['q_h_hru'], self.consumer_data['q_h_0_hr'], U))), 
                     "Гкал/ч"),

                    #Q(h,T) = 1,16 × q(h,T) × (t(h) – t(с)) + Q(ht) 
                    #q(h,T) = qh_u,m × Ui/(1000 × T)

                    ("Расход тепла на ГВС средний (Q(h,T))", 
                     1.16 * (self.consumer_data['q_h'] * U/(1000 * T)) * (t_h - t_c) + 0.4 * (0.005 * self.consumer_data['q_h_0_hr'] * self._calculate_alpha_h(self.consumer_data['q_h_hru'], self.consumer_data['q_h_0_hr'], U)), 
                     "кВт"),

                    ("Расход тепла на ГВС средний (Q(h,T))", 
                    0.0008598452*(1.16 * (self.consumer_data['q_h'] * U/(1000 * T)) * (t_h - t_c) + 0.4 * (0.005 * self.consumer_data['q_h_0_hr'] * self._calculate_alpha_h(self.consumer_data['q_h_hru'], self.consumer_data['q_h_0_hr'], U))), 
                     "Гкал/ч"),
                ]
                
                for param, value, unit in results:
                    self.tree.insert("", "end", values=(param, f"{value:.4f}", unit))
                    
            else:
                messagebox.showerror("Ошибка", "Нет данных о потребителе")
                
        except ValueError as e:
            messagebox.showerror("Ошибка", f"Некорректный ввод: {str(e)}")

    @staticmethod
    def _interpolate(x: float, x_values: List[float], y_values: List[float]) -> float:
        try:
            if x < x_values[0]:
                return LoadCalculator._linear_interpolation(x, x_values[0], x_values[1], y_values[0], y_values[1])
            if x > x_values[-1]:
                return LoadCalculator._linear_interpolation(x, x_values[-2], x_values[-1], y_values[-2], y_values[-1])
            
            for i in range(len(x_values) - 1):
                if x_values[i] <= x <= x_values[i + 1]:
                    return LoadCalculator._linear_interpolation(x, x_values[i], x_values[i + 1], y_values[i], y_values[i + 1])
        except IndexError:
            pass
            
        raise ValueError(LoadCalculator.TEXT["interpolation_error"])

    @staticmethod
    def _linear_interpolation(x: float, x0: float, x1: float, y0: float, y1: float) -> float:
        return y0 + (y1 - y0) * (x - x0) / (x1 - x0)

    #секундная альфа
    def _calculate_alpha(self, q_hru, q_0, U):
        x = (q_hru * U) / (3600 * q_0)
        return ConsumerCalculator._interpolate(x, data.x_values, data.y_values)
    
    #часовая альфа
    def _calculate_alpha_h(self, q_hru, q_0, U):
        x = (q_hru * U) / (q_0)
#        return x
        return LoadCalculator._interpolate(x, data.x_values, data.y_values)
    
    def save_results(self):
        """Сохраняет результаты расчётов в файл (DOCX, CSV или Excel)"""
        if not hasattr(self, 'tree') or not self.tree.get_children():
            messagebox.showerror("Ошибка", "Нет данных для сохранения")
            return

        # Запрашиваем у пользователя путь и тип файла
        file_path = tkinter.filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[(desc, ext) for desc, ext in [ft.value for ft in FileType]],
            title="Сохранить результаты расчёта"
        )
    
        if not file_path:  # Пользователь отменил сохранение
            return

        try:
            # Собираем данные из Treeview
            data = []
            for item in self.tree.get_children():
                values = self.tree.item(item, 'values')
                data.append({
                    "parameter": values[0],
                    "value": values[1],
                    "unit": values[2]
                })

            # Определяем расширение файла
            extension = file_path.split('.')[-1].lower()

            if extension == 'docx':
                self._save_to_docx(file_path, data)
            elif extension == 'csv':
                self._save_to_csv(file_path, data)
            elif extension == 'xlsx':
                self._save_to_excel(file_path, data)
            else:
                messagebox.showerror("Ошибка", "Неподдерживаемый формат файла")
                return

            messagebox.showinfo("Успех", f"Файл успешно сохранён:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")

    def _save_to_docx(self, file_path, data):
        """Сохраняет результаты в документ Word"""
        doc = Document()
    
        # Добавляем заголовок
        doc.add_heading('Результаты расчёта нагрузок', level=1)
    
        # Добавляем таблицу
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
    
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'
        hdr_cells[2].text = 'Ед. изм.'
    
        # Данные таблицы
        for item in data:
            row_cells = table.add_row().cells
            row_cells[0].text = item['parameter']
            row_cells[1].text = item['value']
            row_cells[2].text = item['unit']
    
        doc.save(file_path)

    def _save_to_csv(self, file_path, data):
        """Сохраняет результаты в CSV файл"""
        with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['Параметр', 'Значение', 'Ед. изм.']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        
            writer.writeheader()
            for item in data:
                writer.writerow({
                    'Параметр': item['parameter'],
                    'Значение': item['value'],
                    'Ед. изм.': item['unit']
                })

    def _save_to_excel(self, file_path, data):
        """Сохраняет результаты в Excel файл"""
        wb = Workbook()
        ws = wb.active
        ws.title = "Результаты расчёта"
    
        # Заголовки
        ws.append(['Параметр', 'Значение', 'Ед. изм.'])
    
        # Данные
        for item in data:
            ws.append([item['parameter'], item['value'], item['unit']])
    
        # Автоподбор ширины столбцов
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            ws.column_dimensions[column_letter].width = adjusted_width
    
        wb.save(file_path)

class MainApplication:
    def __init__(self, root):
        self.root = root
        self.root.title("Гидравлический расчёт СП 30.13330.2020")
        self.system_type = tk.StringVar(value="cold")
        
        self._create_interface()
    
    def _create_interface(self):
        # Фрейм для управления
        control_frame = ttk.Frame(self.root, padding="10")
        control_frame.pack(fill=tk.X)
        
        # Фрейм для выбора системы
        system_frame = ttk.Frame(control_frame)
        system_frame.pack(side=tk.LEFT, padx=10)
        
        ttk.Label(system_frame, text="Тип системы:").pack(side=tk.LEFT)
        ttk.Radiobutton(system_frame, text="ХВС", variable=self.system_type, value="cold").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(system_frame, text="ГВС", variable=self.system_type, value="hot").pack(side=tk.LEFT)
        
        # Кнопки
        ttk.Button(control_frame, text="⏬ Импорт из Excel", command=self.import_from_excel).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_frame, text="⏩ Запустить расчёт", command=self.launch_calculator).pack(side=tk.LEFT)
    
    def import_from_excel(self):
        file_path = tk.filedialog.askopenfilename(
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
            title="Выберите файл Excel с данными"
        )
        
        if not file_path:
            return
            
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path)
            sheet = wb.active
            
            # Формат таблицы:
            # Колонка A - номера типа потребителя
            # Колонка B - значения U (количество приборов)
            # Колонка C - диаметры D
            
            data = []
            for row in sheet.iter_rows(min_row=2, values_only=True):  # пропускаем заголовок
                if row[0] and row[1] and row[2]:  # если все три значения есть
                    data.append({
                        't': row[0],  # номер типа потребителя
                        'U': row[1],  # количество приборов
                        'D': row[2]   # диаметр
                    })
            
            # Сохраняем данные для передачи в калькулятор
            self.imported_data = data
            messagebox.showinfo("Успех", f"Успешно загружено {len(data)} участков")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить файл:\n{str(e)}")
    
    def launch_calculator(self):
        if self.system_type.get() == "cold":
            calculator = ConsumerCalculator(tk.Toplevel(self.root))
        else:
            calculator = HotWaterCalculator(tk.Toplevel(self.root))
        
        # Если есть импортированные данные, передаем их в калькулятор
        if hasattr(self, 'imported_data'):
            calculator.load_imported_data(self.imported_data)
        
        calculator.master.eval('tk::Placewindow . center')